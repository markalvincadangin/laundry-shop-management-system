"""
Generate the complete Faith Laundry Shop Management System documentation
as a properly formatted Microsoft Word (.docx) file.
"""

import re
import os
import glob
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Constants ────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS_DIR = os.path.join(BASE_DIR, "diagrams", "png")
OUTPUT_FILE = os.path.join(BASE_DIR, "Faith_Laundry_Shop_System_Documentation.docx")

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(12)
H2_SIZE = Pt(12)
H3_SIZE = Pt(12)
TABLE_SIZE = Pt(10)

TEAM_NAME = "HIMOTECH"
SYSTEM_TITLE = "FAITH LAUNDRY SHOP MANAGEMENT SYSTEM"
SCHOOL = "West Visayas State University"
SCHOOL_ADDR = "La Paz, Iloilo City"
COURSE_LINE = "In Partial Fulfillment of the Requirements for the Systems Analysis and Design Course"

MEMBERS = [
    "Brillantes, Luisa Rose",
    "Cadangin, Mark Alvin",
    "Calisa, Eliza May",
    "De la Cruz, Christian Paul",
    "Serra, Alyanna Bianca",
    "Tacleon, Ellen Mae",
]

# Diagram mappings: which diagrams go with which .md file
# key = md filename, value = list of (png_filename, caption, figure_number)
DIAGRAM_MAP = {
    "06-wbs.md": [
        ("WBS.png", "Work Breakdown Structure", None),
    ],
    "07-project-schedule.md": [
        ("PERT_CPM.png", "PERT/CPM Network Diagram", None),
    ],
    "11-user-stories-and-use-cases.md": [
        ("UCD.png", "Use Case Diagram", None),
    ],
    "16-system-architecture.md": [
        ("SYSTEM_ARCHITECTURE.png", "System Architecture Diagram", None),
    ],
    "17-fdd.md": [
        ("FDD.png", "Functional Decomposition Diagram", None),
    ],
    "18-dfd.md": [
        ("DFD-0.png", "Data Flow Diagram - Level 0 (Context Diagram)", None),
        ("DFD-1.png", "Data Flow Diagram - Level 1", None),
    ],
    "19-erd.md": [
        ("ERD.png", "Entity Relationship Diagram", None),
    ],
}

PHASES = [
    {
        "num": "1",
        "title": "Planning and Initiation",
        "files": [
            ("01-case-study.md", "Case Study"),
            ("02-system-service-request.md", "System Service Request"),
            ("03-project-scope-and-limitations.md", "Project Scope and Limitations"),
        ],
    },
    {
        "num": "2",
        "title": "Phase 2: Project Planning and Management",
        "files": [
            ("04-stakeholder-analysis.md", "Stakeholder Analysis"),
            ("05-roles-and-responsibilities.md", "Roles and Responsibilities"),
            ("06-wbs.md", "Work Breakdown Structure"),
            ("07-project-schedule.md", "Project Schedule"),
            ("08-risk-management-plan.md", "Risk Management Plan"),
        ],
    },
    {
        "num": "3",
        "title": "Phase 3: Systems Analysis and Requirements",
        "files": [
            ("09-process-matrix.md", "Process Matrix"),
            ("10-business-rules.md", "Business Rules"),
            ("11-user-stories-and-use-cases.md", "User Stories and Use Cases"),
            ("12-functional-requirements-checklist.md", "Functional Requirements Checklist"),
            ("13-functional-requirements-matrix.md", "Functional Requirements Matrix"),
            ("14-non-functional-requirements.md", "Non-Functional Requirements"),
        ],
    },
    {
        "num": "4",
        "title": "System Design",
        "files": [
            ("15-technology-stack.md", "Technology Stack"),
            ("16-system-architecture.md", "System Architecture"),
            ("17-fdd.md", "Functional Decomposition Diagram"),
            ("18-dfd.md", "Data Flow Diagram"),
            ("19-erd.md", "Entity Relationship Diagram"),
            ("20-user-interface-design.md", "User Interface Design"),
        ],
    },
    {
        "num": "5",
        "title": "Testing",
        "files": [
            ("21-test-plan.md", "Test Plan and Test Cases"),
        ],
    },
    {
        "num": "6",
        "title": "Deployment And User Manual",
        "files": [
            ("22-user-manual.md", "User Manual"),
            ("23-technical-setup-guide.md", "Technical Setup Guide"),
        ],
    },
]



# ── Helper functions ─────────────────────────────────────────────────────────

def set_run_font(run, size=None, bold=False, italic=False, name=None, color=None):
    """Apply specific font formatting to a run only if explicitly provided."""
    if name is not None:
        run.font.name = name
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = bold
    if italic:
        run.font.italic = italic
    if color:
        run.font.color.rgb = color
        
    # Only apply eastAsia font if name is explicitly provided
    if name is not None:
        r = run._element
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
            r.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc, text, size=None, bold=False, italic=False,
                  alignment=None, space_after=None,
                  space_before=None, first_line_indent=None,
                  line_spacing=None, keep_with_next=False):
    """Add a formatted paragraph. Defaults to native MS Word 'Normal' style without overrides."""
    p = doc.add_paragraph()
    
    if alignment is not None:
        p.alignment = alignment
        
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = space_after
    if space_before is not None:
        pf.space_before = space_before
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if keep_with_next:
        pf.keep_with_next = keep_with_next
        
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, name=FONT_NAME)
            
    return p


def add_centered_text(doc, text, size=None, bold=False, italic=False,
                      space_after=None, space_before=None):
    """Add centered text (used mostly for Title Pages where overrides are needed)."""
    return add_paragraph(
        doc, text, size=size, bold=bold, italic=italic,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Inches(0), space_after=space_after,
        space_before=space_before, line_spacing=1.0,
    )


def add_phase_heading(doc, phase_num, phase_title):
    """Add PHASE X + TITLE as Heading 1."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.style = "Heading 1"
    
    run = p.add_run(phase_title)
    run.font.name = FONT_NAME

def add_section_heading(doc, chapter_idx, section_idx, title):
    """Add Heading 2."""
    title_clean = re.sub(r"^(\d+\.)*\d*\s*", "", title)
    p = doc.add_paragraph(title_clean)
    p.style = "Heading 2"
    for r in p.runs:
        r.font.name = FONT_NAME

def add_subsection_heading(doc, prefix, title):
    """Add Heading 3."""
    title_clean = re.sub(r"^(\d+\.)*\d*\s*", "", title)
    p = doc.add_paragraph(title_clean)
    p.style = "Heading 3"
    for r in p.runs:
        r.font.name = FONT_NAME


def add_figure_placeholder(doc, caption, figure_num, svg_filename=None):
    """Add a diagram image or placeholder box with a caption."""
    svg_path = os.path.join(DIAGRAMS_DIR, svg_filename) if svg_filename else None
    # Also check ui/ subdirectory for UI screenshots
    ui_dir = os.path.join(BASE_DIR, "ui")
    ui_path = os.path.join(ui_dir, svg_filename) if svg_filename else None
    has_file = (svg_path and os.path.exists(svg_path)) or (ui_path and os.path.exists(ui_path))

    # Caption (APA 7th requires Figure captions ABOVE the image)
    # Line 1: Figure X (Bold)
    caption_p1 = doc.add_paragraph(f"Figure {figure_num}")
    caption_p1.style = "Caption"
    if caption_p1.runs:
        caption_p1.runs[0].bold = True
        caption_p1.runs[0].italic = False
        caption_p1.runs[0].font.name = FONT_NAME
    
    # Line 2: Caption Text (Italic)
    caption_p2 = doc.add_paragraph(caption)
    caption_p2.style = "Caption"
    if caption_p2.runs:
        caption_p2.runs[0].italic = True
        caption_p2.runs[0].font.name = FONT_NAME

    if has_file:
        # Resolve the actual image path - check ui/ dir first, then diagrams/png/
        resolved_png = None
        if ui_path and os.path.exists(ui_path):
            resolved_png = ui_path
        elif svg_path:
            png_path = svg_path.replace(".svg", ".png")
            if os.path.exists(png_path):
                resolved_png = png_path
            elif os.path.exists(svg_path):
                resolved_png = None  # SVG exists but no PNG

        if resolved_png:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.first_line_indent = Inches(0)
            run = p.add_run()
            run.add_picture(resolved_png, width=Inches(6.0))
        else:
            # SVG exists but no PNG - add reference placeholder
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.first_line_indent = Inches(0)

            # Create a visible placeholder box
            run = p.add_run(f"[Insert diagram from: diagrams/{svg_filename}]")
            set_run_font(run, size=BODY_SIZE, italic=True, color=RGBColor(0x66, 0x66, 0x66))
    else:
        # MISSING diagram - add placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Inches(0)

        run = p.add_run(f"[PLACEHOLDER: {caption} - To be created]")
        set_run_font(run, size=BODY_SIZE, bold=True, italic=True, color=RGBColor(0xCC, 0x00, 0x00))


def strip_document_control(lines):
    """Remove the Document Control header block from the top of the file."""
    separator_count = 0
    start_idx = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped == "---":
            separator_count += 1
            if separator_count >= 2:
                start_idx = i + 1
                break
    while start_idx < len(lines) and lines[start_idx].strip() in ("", "---"):
        start_idx += 1
    return lines[start_idx:]


def parse_md_content(filepath):
    """Read a markdown file, strip Document Control, and return structured content."""
    with open(filepath, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()

    lines = strip_document_control(raw_lines)
    content = []
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.rstrip("\r\n")
        clean = stripped.strip()

        if clean == "" or clean == "---":
            if in_table and table_rows:
                content.append(("table", table_rows))
                table_rows = []
                in_table = False
            continue

        if clean.startswith(">"):
            continue

        # Table rows
        if "|" in clean and clean.startswith("|"):
            if re.match(r"^\|[\s\-:|]+\|$", clean):
                in_table = True
                continue
            cells = [c.strip() for c in clean.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
                in_table = True
            continue
        else:
            if in_table and table_rows:
                content.append(("table", table_rows))
                table_rows = []
                in_table = False

        # Headings
        if clean.startswith("### "):
            content.append(("h3", clean[4:].strip()))
        elif clean.startswith("## "):
            heading_text = clean[3:].strip()
            if heading_text == "Faith Laundry Shop Management System":
                continue
            content.append(("h2", heading_text))
        elif clean.startswith("# "):
            heading_text = clean[2:].strip()
            if heading_text == "Faith Laundry Shop Management System":
                continue
            content.append(("h1", heading_text))
        elif clean.startswith("- ") or clean.startswith("* "):
            bullet_text = clean[2:].strip()
            content.append(("bullet", bullet_text))
        elif re.match(r"^\d+\.\s", clean):
            num_text = re.sub(r"^\d+\.\s*", "", clean)
            content.append(("numbered", num_text))
        elif clean.startswith("!["):
            match = re.match(r"^!\[(.*?)\]\((.*?)\)", clean)
            if match:
                caption = match.group(1)
                filepath = match.group(2)
                filename = os.path.basename(filepath)
                # Word requires PNGs, but markdown links to SVG
                if filename.endswith('.svg'):
                    filename = filename[:-4] + '.png'
                content.append(("image", (filename, caption)))
            else:
                content.append(("body", clean))
        else:
            content.append(("body", clean))

    if table_rows:
        content.append(("table", table_rows))

    return content


def clean_md_formatting(text):
    """Remove markdown bold/italic markers for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    text = text.replace("<br>", "\n")
    return text


def add_md_table(doc, rows):
    """Add a formatted table from parsed markdown table rows."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)
                pf.line_spacing = 1.0
                pf.first_line_indent = Inches(0)
                run = p.add_run(clean_md_formatting(cell_text))
                set_run_font(run, size=TABLE_SIZE, bold=(i == 0))

    add_paragraph(doc, "", space_after=Pt(6), first_line_indent=Inches(0))


def render_section_content(doc, content, subsection_prefix, next_figure_num, next_table_num):
    """Render parsed markdown content into the Word document."""
    sub_counter = 0
    current_section_title = "Data Table"

    for item_type, data in content:
        if item_type == "h1":
            continue
        elif item_type == "h2":
            sub_counter += 1
            heading_text = re.sub(r"^\d+\.?\s*", "", data)
            current_section_title = heading_text
            add_subsection_heading(doc, f"{subsection_prefix}.{sub_counter}", heading_text)
        elif item_type == "h3":
            heading_text = re.sub(r"^(\d+\.)*\d*\s*", "", data)
            current_section_title = heading_text
            p = doc.add_paragraph(heading_text)
            p.style = "Heading 4"
            if "Heading 4" in doc.styles:
                doc.styles["Heading 4"].font.name = FONT_NAME
            for r in p.runs:
                r.font.name = FONT_NAME
        elif item_type == "body":
            text = clean_md_formatting(data)
            # Normal paragraph - APA 7th requires 0.5 inch first-line indent
            add_paragraph(doc, text, first_line_indent=Inches(0.5))
        elif item_type == "bullet":
            text = clean_md_formatting(data)
            p = doc.add_paragraph(text, style="List Bullet")
            for r in p.runs:
                r.font.name = FONT_NAME
        elif item_type == "numbered":
            text = clean_md_formatting(data)
            p = doc.add_paragraph(text, style="List Number")
            for r in p.runs:
                r.font.name = FONT_NAME
        elif item_type == "image":
            filename, caption = data
            fig_num = next_figure_num()
            add_figure_placeholder(doc, caption, fig_num, filename)
        elif item_type == "table":
            table_num = next_table_num()
            
            # Line 1: Table X (Bold)
            caption_p1 = doc.add_paragraph(f"Table {table_num}")
            caption_p1.style = "Caption"
            if caption_p1.runs:
                caption_p1.runs[0].bold = True
                caption_p1.runs[0].italic = False
                caption_p1.runs[0].font.name = FONT_NAME
                
            # Line 2: Title (Italic)
            caption_p2 = doc.add_paragraph(f"{current_section_title}")
            caption_p2.style = "Caption"
            if caption_p2.runs:
                caption_p2.runs[0].italic = True
                caption_p2.runs[0].font.name = FONT_NAME
                
            add_md_table(doc, data)


# ── Main document builder ───────────────────────────────────────────────────

def build_document():

    doc = Document()
    figure_counter = [0]  # Use list for mutability in nested function
    table_counter = [0]

    def next_figure_num():
        figure_counter[0] += 1
        return figure_counter[0]

    def next_table_num():
        table_counter[0] += 1
        return table_counter[0]

    # ── Page setup ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── Set default font ─────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.first_line_indent = Inches(0.5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Set Headings ─────────────────────────────────────────────────────
    from docx.shared import RGBColor
    
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = H1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 2.0
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = H2_SIZE
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.line_spacing = 2.0
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT_NAME
    h3.font.size = H3_SIZE
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.line_spacing = 2.0
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.keep_with_next = True

    # Check if Heading 4 exists before modifying (to avoid KeyError on some templates)
    try:
        h4 = doc.styles["Heading 4"]
        h4.font.name = FONT_NAME
        h4.font.size = H3_SIZE
        h4.font.bold = True
        h4.font.color.rgb = RGBColor(0, 0, 0)
        h4.paragraph_format.line_spacing = 2.0
        h4.paragraph_format.space_before = Pt(0)
        h4.paragraph_format.space_after = Pt(0)
        h4.paragraph_format.left_indent = Inches(0.0)
        h4.paragraph_format.keep_with_next = True
    except KeyError:
        pass

    # Caption Style
    caption_style = doc.styles["Caption"]
    caption_style.font.name = FONT_NAME
    caption_style.font.size = BODY_SIZE
    caption_style.font.color.rgb = RGBColor(0, 0, 0)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_style.paragraph_format.line_spacing = 2.0
    caption_style.paragraph_format.space_before = Pt(0)
    caption_style.paragraph_format.space_after = Pt(0)

    # Configure List Bullet and List Number globally to inherit Normal but override indent
    for lst_style_name in ["List Bullet", "List Number", "List Paragraph"]:
        try:
            lst_style = doc.styles[lst_style_name]
            lst_style.font.name = FONT_NAME
            lst_style.font.size = BODY_SIZE
            lst_style.paragraph_format.line_spacing = 2.0
            lst_style.paragraph_format.space_after = Pt(0)
            lst_style.paragraph_format.space_before = Pt(0)
            lst_style.paragraph_format.left_indent = Inches(0.5)
            lst_style.paragraph_format.first_line_indent = Inches(-0.25)
            lst_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except KeyError:
            pass

    # ── TITLE PAGE ───────────────────────────────────────────────────────
    for _ in range(3):
        add_centered_text(doc, "")

    add_centered_text(doc, TEAM_NAME, size=Pt(20), bold=True, space_after=Pt(6))

    for _ in range(2):
        add_centered_text(doc, "")

    add_centered_text(doc, SYSTEM_TITLE, size=Pt(18), bold=True, space_after=Pt(24))

    add_centered_text(doc, "System Documentation", size=BODY_SIZE, space_after=Pt(6))
    add_centered_text(doc, "Systems Analysis and Design", size=BODY_SIZE, space_after=Pt(12))
    add_centered_text(doc, SCHOOL, size=BODY_SIZE, space_after=Pt(4))
    add_centered_text(doc, SCHOOL_ADDR, size=BODY_SIZE, space_after=Pt(24))

    add_centered_text(doc, COURSE_LINE, size=BODY_SIZE, italic=True, space_after=Pt(24))

    add_centered_text(doc, "Submitted by:", size=BODY_SIZE, bold=True, space_after=Pt(12))
    for member in MEMBERS:
        add_centered_text(doc, member, size=BODY_SIZE, space_after=Pt(4))

    add_centered_text(doc, "", space_after=Pt(12))
    add_centered_text(doc, TEAM_NAME, size=Pt(14), bold=True, space_after=Pt(24))
    add_centered_text(doc, "May 2026", size=BODY_SIZE)

    # ── ABSTRACT ─────────────────────────────────────────────────────────
    doc.add_page_break()
    add_centered_text(doc, "ABSTRACT", size=H1_SIZE, bold=True, space_after=Pt(24))

    abstract_paragraphs = [
        "Faith Laundry Shop, located in Ilaya, Tabuc Suba, Jaro, Iloilo City, is a small-scale laundry service that has been operating since 2022. The business currently relies on manual processes including handwritten logbooks, physical tags, and paper-based receipts for managing daily operations. These methods have led to recurring issues such as time-consuming record keeping, occasional order mix-ups during peak hours, limited order tracking capability, and the absence of automated sales reporting.",
        "This system documentation presents the analysis, design, and development of the Faith Laundry Shop Management System \u2014 a web-based application intended to digitize and streamline the shop\u2019s core operations. The system supports automated computation of laundry charges based on a per-load pricing structure, real-time order status tracking, digital receipt generation, and automated daily and monthly sales reporting.",
        "The system was developed using a modern technology stack comprising Next.js for the frontend, Spring Boot with Java 21 for the backend, and PostgreSQL for the database. The documentation covers the complete systems development lifecycle, including the case study, requirements specification, system design, implementation, testing, and user documentation.",
    ]
    for para_text in abstract_paragraphs:
        add_paragraph(doc, para_text, first_line_indent=Inches(0.5))

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.line_spacing = 2.0
    pf.first_line_indent = Inches(0.5)
    kw_label = p.add_run("Keywords: ")
    set_run_font(kw_label, bold=True, italic=True)
    kw_text = p.add_run("laundry management system, web application, order tracking, automated reporting, systems analysis and design")
    set_run_font(kw_text, italic=True)

    # ── TABLE OF CONTENTS PLACEHOLDER ────────────────────────────────────
    doc.add_page_break()
    add_centered_text(doc, "TABLE OF CONTENTS", size=H1_SIZE, bold=True, space_after=Pt(24))
    add_paragraph(
        doc,
        "[To generate: In Word, go to References > Table of Contents > Automatic Table 1. "
        "This will auto-populate based on your Heading styles.]",
        italic=True, first_line_indent=Inches(0),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── LIST OF FIGURES PLACEHOLDER ──────────────────────────────────────
    doc.add_page_break()
    add_centered_text(doc, "LIST OF FIGURES", size=H1_SIZE, bold=True, space_after=Pt(24))
    add_paragraph(
        doc,
        "[To generate: In Word, go to References > Insert Table of Figures.]",
        italic=True, first_line_indent=Inches(0),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── LIST OF TABLES PLACEHOLDER ───────────────────────────────────────
    doc.add_page_break()
    add_centered_text(doc, "LIST OF TABLES", size=H1_SIZE, bold=True, space_after=Pt(24))
    add_paragraph(
        doc,
        "[To generate: In Word, go to References > Insert Table of Figures > select 'Table' label.]",
        italic=True, first_line_indent=Inches(0),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── PHASES ─────────────────────────────────────────────────────────
    for ch_idx, phase in enumerate(PHASES, start=1):
        add_phase_heading(doc, phase["num"], phase["title"])

        for sec_idx, (filename, section_title) in enumerate(phase["files"], start=1):
            filepath = os.path.join(BASE_DIR, filename)

            add_section_heading(doc, ch_idx, sec_idx, section_title)

            if not os.path.exists(filepath):
                add_paragraph(doc, f"[Content pending: {filename}]", italic=True)
                continue

            # Parse and render text content
            content = parse_md_content(filepath)
            subsection_prefix = f"{ch_idx}.{sec_idx}"
            render_section_content(doc, content, subsection_prefix, next_figure_num, next_table_num)

            # Render mapped diagrams
            if filename in DIAGRAM_MAP:
                for diag_file, diag_caption, _ in DIAGRAM_MAP[filename]:
                    add_figure_placeholder(doc, diag_caption, next_figure_num(), diag_file)

    # ── REFERENCES ───────────────────────────────────────────
    doc.add_page_break()
    p_ref_title = doc.add_paragraph()
    p_ref_title.style = "Heading 1"
    run_ref = p_ref_title.add_run("References")
    run_ref.font.name = FONT_NAME

    references = [
        "GeeksforGeeks. (2024). FDD full form. https://www.geeksforgeeks.org/software-engineering/fdd-full-form/",
        "Meta Platforms. (2024). React [Computer software]. https://react.dev/",
        "Oracle. (2023). Java (Version 21) [Computer software]. https://www.oracle.com/java/",
        "PostgreSQL Global Development Group. (2024). PostgreSQL [Computer software]. https://www.postgresql.org/",
        "Tailwind Labs. (2024). Tailwind CSS [Computer software]. https://tailwindcss.com/",
        "Valacich, J. S., & George, J. F. (2020). Modern systems analysis and design (9th ed.). Pearson.",
        "Vercel. (2024). Next.js [Computer software]. https://nextjs.org/",
        "VMware. (2024). Spring Boot [Computer software]. https://spring.io/projects/spring-boot"
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = 2.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # APA 7th requires a 0.5 inch hanging indent for references
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = BODY_SIZE
            # Automatically italicize the titles based on APA rules (titles come after the date)
            if "React" in r.text:
                r.text = r.text.replace("React", "")
                i_run = p.add_run("React")
                set_run_font(i_run, italic=True)
                p.add_run(" [Computer software]. https://react.dev/")
            elif "Java (Version 21)" in r.text:
                r.text = "Oracle. (2023). "
                i_run = p.add_run("Java")
                set_run_font(i_run, italic=True)
                p.add_run(" (Version 21) [Computer software]. https://www.oracle.com/java/")
            elif "FDD full form" in r.text:
                r.text = "GeeksforGeeks. (2024). "
                i_run = p.add_run("FDD full form")
                set_run_font(i_run, italic=True)
                p.add_run(". https://www.geeksforgeeks.org/software-engineering/fdd-full-form/")
            elif "PostgreSQL [Computer" in r.text:
                r.text = "PostgreSQL Global Development Group. (2024). "
                i_run = p.add_run("PostgreSQL")
                set_run_font(i_run, italic=True)
                p.add_run(" [Computer software]. https://www.postgresql.org/")
            elif "Tailwind CSS [Computer" in r.text:
                r.text = "Tailwind Labs. (2024). "
                i_run = p.add_run("Tailwind CSS")
                set_run_font(i_run, italic=True)
                p.add_run(" [Computer software]. https://tailwindcss.com/")
            elif "Modern systems analysis and design" in r.text:
                r.text = "Valacich, J. S., & George, J. F. (2020). "
                i_run = p.add_run("Modern systems analysis and design")
                set_run_font(i_run, italic=True)
                p.add_run(" (9th ed.). Pearson.")
            elif "Next.js" in r.text:
                r.text = "Vercel. (2024). "
                i_run = p.add_run("Next.js")
                set_run_font(i_run, italic=True)
                p.add_run(" [Computer software]. https://nextjs.org/")
            elif "Spring Boot" in r.text:
                r.text = "VMware. (2024). "
                i_run = p.add_run("Spring Boot")
                set_run_font(i_run, italic=True)
                p.add_run(" [Computer software]. https://spring.io/projects/spring-boot")

    # ── APPENDICES PLACEHOLDER ───────────────────────────────────────────
    doc.add_page_break()
    add_centered_text(doc, "APPENDICES", size=H1_SIZE, bold=True, space_after=Pt(24))

    appendix_items = [
        "Appendix A: Client Interview Transcript",
        "Appendix B: Source Code Listing",
        "Appendix C: System Screenshots",
        "Appendix D: Gantt Chart",
    ]
    for item in appendix_items:
        add_paragraph(
            doc, f"[{item} - To be added]",
            italic=True, first_line_indent=Inches(0),
            space_after=Pt(12),
        )

    # ── SAVE ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"\n[OK] Document generated successfully!")
    print(f"   File: {OUTPUT_FILE}")
    print(f"   Phases: {len(PHASES)}")
    total_sections = sum(len(ch['files']) for ch in PHASES)
    print(f"   Sections: {total_sections}")
    print(f"   Figures: {figure_counter[0]} (including placeholders)")


if __name__ == "__main__":
    build_document()
